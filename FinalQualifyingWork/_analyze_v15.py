#!/usr/bin/env python3
"""Анализ БояриновИР_Дипломная v1.5.docx"""
from __future__ import annotations

import re
from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

HERE = Path(__file__).parent
DOC = HERE / "БояриновИР_Дипломная v1.5.docx"
OUT = HERE / "_analysis_v15.txt"
MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def main() -> None:
    doc = Document(DOC)
    lines: list[str] = []

    paras = doc.paragraphs
    texts = [p.text.strip() for p in paras]
    nonempty = [t for t in texts if t]

    # headings by style
    headings: list[tuple[str, str]] = []
    for p in paras:
        name = (p.style.name if p.style else "") or ""
        t = p.text.strip()
        if not t:
            continue
        if name.startswith("Heading") or re.match(r"^\d+(\.\d+)*\s", t):
            headings.append((name, t[:120]))

    # figures / tables captions
    fig_caps = [t for t in texts if re.match(r"^Рисунок\s+\d+", t, re.I)]
    tab_caps = [t for t in texts if re.match(r"^Таблица\s+\d+", t, re.I)]

    # formulas
    display_math = sum(
        1 for p in paras if p._element.findall(f".//{{{MATH_NS}}}oMathPara")
    )
    inline_math_paras = []
    legend_paras = []
    for p in paras:
        el = p._element
        has_para = bool(el.findall(f".//{{{MATH_NS}}}oMathPara"))
        inline = el.findall(f".//{{{MATH_NS}}}oMath")
        if has_para:
            continue
        if inline and texts[paras.index(p)].startswith("где"):
            legend_paras.append(p)
        elif inline:
            inline_math_paras.append(p)

    # cross-refs patterns
    ref_рис = len(re.findall(r"рис\.\s*\d+", " ".join(texts), re.I))
    ref_табл = len(re.findall(r"табл\.\s*\d+", " ".join(texts), re.I))
    ref_см = len(re.findall(r"см\.\s*\(\d", " ".join(texts), re.I))
    ref_прил = len(re.findall(r"прил\.\s*[А-ЯA-Z]", " ".join(texts), re.I))
    ref_формула = len(re.findall(r"\(3\.\d\)|\(4\.\d\)", " ".join(texts)))

    # experiments
    exp_heads = [t for t in texts if re.match(r"^Эксперимент\s+\d+", t, re.I)]

    # placeholders / stubs
    placeholders = [t for t in texts if "PLACEHOLDER" in t.upper() or "ЗАГЛУШК" in t.upper() or "БУДЕТ ДОБАВЛЕН" in t.upper()]

    # SEQ fields (caption auto-number)
    seq_fields = 0
    for p in paras:
        for el in p._element.iter():
            if el.tag == qn("w:instrText") and "SEQ" in (el.text or ""):
                seq_fields += 1

    # word count rough
    words = sum(len(t.split()) for t in nonempty)

    # chapters from headings
    chapters = [t for _, t in headings if re.match(r"^[1-4]\s", t) or re.match(r"^Глава", t, re.I)]

    # formula labels (3.1)..(4.2)
    formula_labels = []
    for p in paras:
        t = p.text.strip()
        if re.fullmatch(r"\(3\.\d\)|\(4\.\d\)", t) or (
            t.endswith(")") and re.search(r"\(3\.\d\)|\(4\.\d\)", t) and "oMathPara" in p._element.xml
        ):
            m = re.search(r"\(3\.\d\)|\(4\.\d\)", t)
            if m:
                formula_labels.append(m.group())

    # duplicate figure numbers in captions
    fig_nums = []
    for c in fig_caps:
        m = re.match(r"^Рисунок\s+(\d+)", c, re.I)
        if m:
            fig_nums.append(int(m.group(1)))
    fig_dup = [n for n, cnt in Counter(fig_nums).items() if cnt > 1]

    tab_nums = []
    for c in tab_caps:
        m = re.match(r"^Таблица\s+(\d+)", c, re.I)
        if m:
            tab_nums.append(int(m.group(1)))
    tab_dup = [n for n, cnt in Counter(tab_nums).items() if cnt > 1]

    # tables in document
    n_tables = len(doc.tables)

    lines.append("=== БояриновИР_Дипломная v1.5.docx ===\n")
    lines.append(f"Абзацев (всего / непустых): {len(paras)} / {len(nonempty)}")
    lines.append(f"Слов (прибл.): {words}")
    lines.append(f"Таблиц Word: {n_tables}")
    lines.append(f"Подписей «Рисунок N»: {len(fig_caps)}")
    lines.append(f"Подписей «Таблица N»: {len(tab_caps)}")
    lines.append(f"Уравнений display (oMathPara): {display_math}")
    lines.append(f"Абзацев «где …» с inline OMML: {len(legend_paras)}")
    lines.append(f"Метки формул в тексте (3.N)/(4.N): {ref_формула}")
    lines.append(f"Отсылки «см. (N)»: {ref_см}")
    lines.append(f"Отсылки «рис.»: {ref_рис}")
    lines.append(f"Отсылки «табл.»: {ref_табл}")
    lines.append(f"Отсылки «прил.»: {ref_прил}")
    lines.append(f"Полей SEQ в подписях: {seq_fields}")
    lines.append(f"Заголовков экспериментов: {len(exp_heads)}")
    if exp_heads:
        lines.append("  " + "; ".join(exp_heads[:14]))
    lines.append(f"Заглушек/placeholder: {len(placeholders)}")
    for ph in placeholders[:8]:
        lines.append(f"  - {ph[:100]}")
    if fig_dup:
        lines.append(f"Дубли номеров рисунков: {fig_dup}")
    if tab_dup:
        lines.append(f"Дубли номеров таблиц: {tab_dup}")

    lines.append("\n--- Заголовки (первые 40) ---")
    for name, t in headings[:40]:
        lines.append(f"[{name}] {t}")

    lines.append("\n--- Подписи рисунков ---")
    for c in fig_caps:
        lines.append(c[:140])

    lines.append("\n--- Формулы (display) ---")
    for p in paras:
        if p._element.findall(f".//{{{MATH_NS}}}oMathPara"):
            idx = paras.index(p)
            prev = texts[max(0, idx - 1)][:80] if idx else ""
            lines.append(f"p{idx}: label={p.text.strip()!r} after: {prev!r}")

    lines.append("\n--- Эксперименты: таблицы в тексте? ---")
    for eh in exp_heads:
        # find paragraph index
        for i, t in enumerate(texts):
            if t == eh:
                chunk = " ".join(texts[i : i + 25])
                has_tab = bool(re.search(r"табл\.\s*\d+", chunk, re.I))
                has_рис = bool(re.search(r"рис\.\s*\d+", chunk, re.I))
                lines.append(f"{eh}: табл={has_tab} рис={has_рис}")
                break

    lines.append("\n--- Потенциальные проблемы ---")
    # figures without ref
    for n in sorted(set(fig_nums)):
        if not re.search(rf"рис\.\s*{n}\b", " ".join(texts), re.I):
            if n <= 50:  # only main body figures
                lines.append(f"Рисунок {n} возможно без отсылки в тексте")

    # missing formulas
    for lab in ["(3.1)", "(3.2)", "(3.3)", "(3.4)", "(3.5)", "(3.6)", "(4.1)", "(4.2)"]:
        if lab not in " ".join(texts):
            lines.append(f"Нет метки {lab}")

    OUT.write_text("\n".join(lines), encoding="utf-8")
    print(OUT.read_text(encoding="utf-8"))


if __name__ == "__main__":
    main()
