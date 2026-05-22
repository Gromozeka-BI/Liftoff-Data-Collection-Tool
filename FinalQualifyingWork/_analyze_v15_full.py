#!/usr/bin/env python3
from __future__ import annotations

import re
from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

DOC = Path(__file__).parent / "БояриновИР_Дипломная v1.5.docx"
OUT = Path(__file__).parent / "_analysis_v15_full.txt"
MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"


def para_has_image(p) -> bool:
    return bool(
        p._element.findall(".//{%s}blip" % A_NS)
        or p._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/picture}pic")
    )


def main() -> None:
    doc = Document(DOC)
    paras = doc.paragraphs
    texts = [p.text for p in paras]
    joined = "\n".join(texts)

    lines: list[str] = ["# Анализ v1.5\n"]

    # объём
    nonempty = [t.strip() for t in texts if t.strip()]
    words = sum(len(t.split()) for t in nonempty)
    lines.append(f"Абзацев: {len(paras)}, непустых: {len(nonempty)}, слов ~{words}")
    lines.append(f"Таблиц Word: {len(doc.tables)}")

    # рисунки: подписи в любом месте абзаца
    fig_pat = re.compile(r"Рисунок\s+(\d+)", re.I)
    fig_caps = []
    for i, t in enumerate(texts):
        if fig_pat.search(t):
            fig_caps.append((i, t.strip()[:150]))
    fig_nums = [int(fig_pat.search(t).group(1)) for _, t in fig_caps]
    lines.append(f"Абзацев с «Рисунок N»: {len(fig_caps)}, уникальных номеров: {len(set(fig_nums))}")
    if fig_nums:
        lines.append(f"  диапазон номеров: {min(fig_nums)}–{max(fig_nums)}")
    dup = [n for n, c in Counter(fig_nums).items() if c > 1]
    if dup:
        lines.append(f"  дубли номеров: {sorted(dup)}")

    tab_pat = re.compile(r"Таблица\s+(\d+)", re.I)
    tab_caps = [(i, t.strip()[:150]) for i, t in enumerate(texts) if tab_pat.search(t)]
    tab_nums = [int(tab_pat.search(t).group(1)) for _, t in tab_caps]
    lines.append(f"Абзацев с «Таблица N»: {len(tab_caps)}, уникальных: {len(set(tab_nums))}")
    if tab_nums:
        lines.append(f"  диапазон: {min(tab_nums)}–{max(tab_nums)}")

    # изображения в абзацах
    img_paras = sum(1 for p in paras if para_has_image(p))
    lines.append(f"Абзацев с встроенным изображением: {img_paras}")

    # формулы
    omath_para = sum(1 for p in paras if p._element.findall(f".//{{{MATH_NS}}}oMathPara"))
    omath_any = sum(1 for p in paras if p._element.findall(f".//{{{MATH_NS}}}oMath"))
    label_only = [i for i, t in enumerate(texts) if re.fullmatch(r"\s*\(3\.\d\)\s*|\s*\(4\.\d\)\s*", t.strip())]
    legend = [i for i, t in enumerate(texts) if t.strip().startswith("где")]
    lines.append(f"oMathPara (display): {omath_para}")
    lines.append(f"Абзацев с oMath (любой): {omath_any}")
    lines.append(f"Строк-меток (3.N)/(4.N) отдельным абзацем: {len(label_only)}")
    lines.append(f"Абзацев «где …»: {len(legend)}")

    # отсылки
    for name, pat in [
        ("см. (N)", r"см\.\s*\([34]\.\d\)"),
        ("рис.", r"рис\.\s*\d+"),
        ("табл.", r"табл\.\s*\d+"),
        ("прил.", r"прил\.\s*[А-ЯA-Z]"),
        ("формула (3/4)", r"\([34]\.\d\)"),
    ]:
        lines.append(f"Отсылок {name}: {len(re.findall(pat, joined, re.I))}")

    # структура: ключевые разделы
    markers = [
        "ТИТУЛ", "РЕФЕРАТ", "СОДЕРЖАНИЕ", "ПЕРЕЧЕНЬ", "ВВЕДЕНИЕ",
        "Глава 1", "1 Анализ", "Глава 2", "Глава 3", "Глава 4",
        "ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗОВАННЫХ", "ПРИЛОЖЕНИЕ",
    ]
    lines.append("\n## Найденные разделы")
    for m in markers:
        hits = [t.strip()[:80] for t in texts if m.lower() in t.lower()[:60]]
        if hits:
            lines.append(f"- {m}: {hits[0]}")

    # эксперименты
    exps = [t.strip() for t in texts if re.match(r"^Эксперимент\s+\d+", t.strip(), re.I)]
    lines.append(f"\n## Эксперименты: {len(exps)}")
    for e in exps:
        lines.append(f"  - {e}")

    # рисунки без отсылки (1–46)
    missing_fig = []
    for n in range(1, 47):
        if any(f"рис. {n}" in t.lower() or f"рис.{n}" in t.lower() for t in texts):
            continue
        if n in fig_nums:
            missing_fig.append(n)
    lines.append(f"\n## Рисунки с подписью, но без «рис. N» в тексте: {missing_fig[:20]}{'...' if len(missing_fig)>20 else ''}")

    # заглушки
    stubs = [t[:100] for t in texts if re.search(r"будет добавлен|заглуш|placeholder|TODO", t, re.I)]
    lines.append(f"\n## Заглушки: {len(stubs)}")
    for s in stubs[:5]:
        lines.append(f"  - {s}")

    # стили заголовков глав
    h1 = [p.text.strip() for p in paras if (p.style and p.style.name == "Heading 1")]
    lines.append(f"\n## Heading 1 ({len(h1)})")
    for h in h1[:15]:
        lines.append(f"  - {h[:100]}")

    # примеры подписей рисунков
    lines.append("\n## Все подписи «Рисунок» (первые 15 и последние 5)")
    for i, t in fig_caps[:15]:
        lines.append(f"  p{i}: {t}")
    for i, t in fig_caps[-5:]:
        lines.append(f"  p{i}: {t}")

    # формулы: контекст
    lines.append("\n## Формулы (3.1)–(4.2): контекст")
    for lab in ["(3.1)", "(3.2)", "(3.3)", "(3.4)", "(3.5)", "(3.6)", "(4.1)", "(4.2)"]:
        for i, t in enumerate(texts):
            if lab in t and t.strip() in (lab, f"\t{lab}"):
                prev = texts[i - 1].strip()[:100] if i else ""
                nxt = texts[i + 1].strip()[:100] if i + 1 < len(texts) else ""
                has_omath = "oMath" in paras[i]._element.xml
                lines.append(f"{lab} p{i}: oMath={has_omath}; prev={prev!r}; next={nxt!r}")
                break

    OUT.write_text("\n".join(lines), encoding="utf-8")
    print(OUT.read_text(encoding="utf-8"))


if __name__ == "__main__":
    main()
