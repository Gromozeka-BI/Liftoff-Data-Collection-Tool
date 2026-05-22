#!/usr/bin/env python3
"""
Копия БояриновИР_Дипломная v1.3.docx → v1.4:
- формулы как уравнения Word (OMML через LaTeX);
- отсылки «см. (3.N)» / «см. (4.N)» в тексте абзацев;
- абзац «где …» с расшифровкой символов под каждой формулой
  (обозначения — inline-уравнения Word, как в формуле).
"""
from __future__ import annotations

import re
import shutil
from pathlib import Path
from typing import Literal

import latex2mathml.converter as latex2mathml
import mathml2omml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.text.paragraph import Paragraph

LegendSeg = tuple[Literal["text", "math"], str]

HERE = Path(__file__).parent
SRC = HERE / "БояриновИР_Дипломная v1.3.docx"
DST = HERE / "БояриновИР_Дипломная v1.4.docx"
CHANGELOG = HERE / "CHANGELOG_v14_формулы.md"

MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

# needle, LaTeX, label, фраза-отсылка, сегменты абзаца «где …» (text | math)
FORMULA_ENTRIES: list[tuple[str, str, str, str, list[LegendSeg]]] = [
    (
        "между текущим наблюдением и узлами референса",
        r"\tilde{u}_k = \frac{u_k - \mu_k}{\sigma_k}, \quad k = 1,\ldots,K",
        "(3.1)",
        " (см. (3.1))",
        [
            ("text", "где "),
            ("math", r"\tilde{u}_k"),
            ("text", " — нормализованный "),
            ("math", r"k"),
            ("text", "-й признак наблюдения; "),
            ("math", r"u_k"),
            ("text", " — сглаженное значение после предобработки; "),
            ("math", r"\mu_k"),
            ("text", ", "),
            ("math", r"\sigma_k"),
            ("text", " — среднее и СКО "),
            ("math", r"k"),
            ("text", "-го признака по эталонному кругу; "),
            ("math", r"K"),
            ("text", " — число признаков (каналов RC)."),
        ],
    ),
    (
        "работать с циклическим кругом",
        r"v_i \leftarrow \mathrm{clip}(v_i + \eta_v, v_{\min}, v_{\max}), "
        r"\quad s_i \leftarrow (s_i + v_i \Delta t + \eta_s) \bmod L",
        "(3.2)",
        " (см. (3.2))",
        [
            ("text", "где "),
            ("math", r"v_i"),
            ("text", " — скорость "),
            ("math", r"i"),
            ("text", "-й частицы вдоль трассы, м/с; "),
            ("math", r"\eta_v"),
            ("text", " — гауссовское возмущение скорости (process_noise_v); "),
            ("math", r"v_{\min}"),
            ("text", ", "),
            ("math", r"v_{\max}"),
            ("text", " — допустимые границы скорости; "),
            ("math", r"s_i"),
            ("text", " — дуговая координата частицы, м; "),
            ("math", r"\Delta t"),
            ("text", " — интервал между наблюдениями, с; "),
            ("math", r"\eta_s"),
            ("text", " — гауссовское возмущение положения (process_noise_s); "),
            ("math", r"L"),
            ("text", " — длина референсного круга, м; операция "),
            ("math", r"\bmod L"),
            ("text", " обеспечивает цикличность трассы."),
        ],
    ),
    (
        "исключать throttle из метрики при переносе между разными дронами",
        r"d_i^2 = \sum_k w_k \left(f_{\mathrm{ref},k}(s_i) - \tilde{u}_k\right)^2, "
        r"\quad w_i \propto w_i \exp\left(-\frac{d_i^2}{2\sigma_{\mathrm{obs}}^2}\right), "
        r"\quad \sum_i w_i = 1",
        "(3.3)",
        " (см. (3.3))",
        [
            ("text", "где "),
            ("math", r"d_i^2"),
            ("text", " — взвешенная квадратичная невязка RC-признаков для частицы "),
            ("math", r"i"),
            ("text", "; "),
            ("math", r"k"),
            ("text", " — индекс признака; "),
            ("math", r"w_k"),
            ("text", " — вес канала в метрике (при channel_weights); "),
            ("math", r"f_{\mathrm{ref},k}(s_i)"),
            ("text", " — "),
            ("math", r"k"),
            ("text", "-й нормализованный признак референса в точке "),
            ("math", r"s_i"),
            ("text", "; "),
            ("math", r"\tilde{u}_k"),
            ("text", " — текущее наблюдение по (3.1); "),
            ("math", r"w_i"),
            ("text", " — вес частицы; "),
            ("math", r"\sigma_{\mathrm{obs}}"),
            ("text", " — параметр obs_sigma; нормировка "),
            ("math", r"\sum_i w_i = 1"),
            ("text", " выполняется после обновления."),
        ],
    ),
    (
        "корректно обрабатывает границу старт/финиш",
        r"\mathrm{ESS} = \frac{1}{\sum_i w_i^2}, \quad "
        r"\theta_i = \frac{2\pi s_i}{L}, \quad "
        r"s_{\mathrm{est}} = \frac{L}{2\pi}\mathrm{atan2}\!\left(\sum_i w_i \sin\theta_i,\, "
        r"\sum_i w_i \cos\theta_i\right)",
        "(3.4)",
        " (см. (3.4))",
        [
            ("text", "где "),
            ("math", r"\mathrm{ESS}"),
            ("text", " — эффективное число частиц; "),
            ("math", r"w_i"),
            ("text", " — вес "),
            ("math", r"i"),
            ("text", "-й частицы; "),
            ("math", r"\theta_i"),
            ("text", " — угол на окружности, соответствующий положению "),
            ("math", r"s_i"),
            ("text", "; "),
            ("math", r"L"),
            ("text", " — длина круга, м; "),
            ("math", r"s_{\mathrm{est}}"),
            ("text", " — оценка дуговой координаты (circular mean); "),
            ("math", r"\mathrm{atan2}"),
            ("text", " — среднее по окружности с учётом перехода старт/финиш."),
        ],
    ),
    (
        "расстояние до камерного наблюдения преобразуется в правдоподобие",
        r"w_i \leftarrow w_i \exp\left(-\frac{\left\|p(s_i) - z_{\mathrm{obs}}\right\|^2}"
        r"{2\sigma_{\mathrm{cam}}^2}\right)",
        "(3.5)",
        " (см. (3.5))",
        [
            ("text", "где "),
            ("math", r"w_i"),
            ("text", " — вес частицы до и после камерного обновления; "),
            ("math", r"p(s_i)"),
            ("text", " — 3D-позиция частицы на референсе; "),
            ("math", r"z_{\mathrm{obs}}"),
            ("text", " — абсолютное камерное наблюдение xyz_obs; "),
            ("math", r"\sigma_{\mathrm{cam}}"),
            ("text", " — оценка неопределённости камеры (sigma_cam); "),
            ("math", r"\|\cdot\|"),
            ("text", " — евклидово расстояние в локальной СК трассы."),
        ],
    ),
    (
        "Внутренний критерий работоспособности был задан как «p90 < 15 м»",
        r"e_t = \left\|p_{\mathrm{gt}}(t) - p_{\mathrm{est}}(t)\right\|, \quad "
        r"\mathrm{p90\_err\_m} = Q_{0.9}(\{e_t\}), \quad "
        r"\mathrm{p95\_err\_m} = Q_{0.95}(\{e_t\})",
        "(3.6)",
        " (см. (3.6))",
        [
            ("text", "где "),
            ("math", r"e_t"),
            ("text", " — модуль ошибки локализации в момент "),
            ("math", r"t"),
            ("text", ", м; "),
            ("math", r"p_{\mathrm{gt}}(t)"),
            ("text", " — эталонная (истинная) позиция; "),
            ("math", r"p_{\mathrm{est}}(t)"),
            ("text", " — оценка фильтра; "),
            ("math", r"\mathrm{p90\_err\_m}"),
            ("text", " и "),
            ("math", r"\mathrm{p95\_err\_m}"),
            ("text", " — 90-й и 95-й процентили множества "),
            ("math", r"\{e_t\}"),
            ("text", "; "),
            ("math", r"Q_{0.9}"),
            ("text", " и "),
            ("math", r"Q_{0.95}"),
            ("text", " — соответствующие квантили."),
        ],
    ),
    (
        "обратно в широту, долготу и высоту",
        r"\mathrm{East} = R\cos\varphi_0\,\Delta\lambda, \quad "
        r"\mathrm{North} = R\,\Delta\varphi, \quad "
        r"\mathrm{Up} = \Delta h, \quad "
        r"\mathbf{x}_{\mathrm{ENU}} = \mathbf{b}_x x_{\mathrm{loc}} + \mathbf{b}_z z_{\mathrm{loc}}"
        r" + \mathbf{r}_{\mathrm{origin}}",
        "(4.1)",
        " (см. (4.1))",
        [
            ("text", "где "),
            ("math", r"\mathrm{East}"),
            ("text", ", "),
            ("math", r"\mathrm{North}"),
            ("text", ", "),
            ("math", r"\mathrm{Up}"),
            ("text", " — компоненты ENU относительно опорной точки; "),
            ("math", r"R"),
            ("text", " — радиус Земли; "),
            ("math", r"\varphi_0"),
            ("text", " — широта опорной точки; "),
            ("math", r"\Delta\lambda"),
            ("text", ", "),
            ("math", r"\Delta\varphi"),
            ("text", ", "),
            ("math", r"\Delta h"),
            ("text", " — приращения долготы, широты и высоты; "),
            ("math", r"x_{\mathrm{loc}}"),
            ("text", ", "),
            ("math", r"z_{\mathrm{loc}}"),
            ("text", " — координаты в локальной СК трассы; "),
            ("math", r"\mathbf{b}_x"),
            ("text", ", "),
            ("math", r"\mathbf{b}_z"),
            ("text", " — единичные векторы осей трассы; "),
            ("math", r"\mathbf{r}_{\mathrm{origin}}"),
            ("text", " — начало локальной системы; "),
            ("math", r"\mathbf{x}_{\mathrm{ENU}}"),
            ("text", " — вектор в ENU."),
        ],
    ),
    (
        "категории HEPU в документации внешней системы задаются как 95% граница",
        r"P\!\left(\|e\|_{\mathrm{hor}} < T_k\right) \geq 0{,}95 "
        r"\quad (\text{оценка по } \mathrm{p95\_err\_m})",
        "(4.2)",
        " (см. (4.2))",
        [
            ("text", "где "),
            ("math", r"P(\cdot)"),
            ("text", " — вероятность попадания горизонтальной ошибки в допуск; "),
            ("math", r"\|e\|_{\mathrm{hor}}"),
            ("text", " — модуль ошибки в горизонтальной плоскости; "),
            ("math", r"T_k"),
            ("text", " — порог "),
            ("math", r"k"),
            ("text", "-й категории HEPU; 0,95 — требуемая доля событий по документации "
            "внешней системы; "),
            ("math", r"\mathrm{p95\_err\_m}"),
            ("text", " — оценка по 95-му процентилю ошибки из экспериментов (формула (3.6))."),
        ],
    ),
]

LEGACY_FORMULA_RE = re.compile(r"^\s*.*\((3\.\d|4\.\d)\)\s*$")


def insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_el = OxmlElement("w:p")
    paragraph._element.addnext(new_el)
    return Paragraph(new_el, paragraph._parent)


def paragraph_index(doc: Document, paragraph: Paragraph) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p._element is paragraph._element:
            return i
    return -1


def next_is_equation_block(paragraph: Paragraph, label: str) -> bool:
    nxt = paragraph._element.getnext()
    if nxt is None:
        return False
    if label in "".join(nxt.itertext()):
        return True
    return bool(nxt.findall(".//{%s}oMath" % MATH_NS))


def find_equation_paragraph(anchor: Paragraph, label: str) -> Paragraph | None:
    nxt = anchor._element.getnext()
    if nxt is None:
        return None
    if label in "".join(nxt.itertext()) and nxt.findall(".//{%s}oMath" % MATH_NS):
        return Paragraph(nxt, anchor._parent)
    return None


def is_legend_paragraph(paragraph: Paragraph) -> bool:
    """Абзац «где …» с inline-формулами (не блок уравнения по центру)."""
    el = paragraph._element
    if el.findall(".//{%s}oMathPara" % MATH_NS):
        return False
    text = paragraph.text.strip()
    if not text.startswith("где"):
        return False
    if el.findall(".//{%s}oMath" % MATH_NS):
        return True
    # старый текстовый вариант «где ũ_k …»
    return " — " in text and len(text) < 600


def next_is_legend_block(equation_paragraph: Paragraph, _label: str) -> bool:
    nxt = equation_paragraph._element.getnext()
    if nxt is None:
        return False
    return is_legend_paragraph(Paragraph(nxt, equation_paragraph._parent))


def latex_to_omml_element(latex: str):
    mathml = latex2mathml.convert(latex)
    omml = mathml2omml.convert(mathml)
    wrapped = f'<root xmlns:m="{MATH_NS}">{omml}</root>'
    root = parse_xml(wrapped)
    omath = root.find(f".//{{{MATH_NS}}}oMath")
    if omath is None:
        raise RuntimeError(f"no oMath for: {latex!r}")
    return omath


def add_inline_math_run(paragraph: Paragraph, latex: str) -> None:
    run = paragraph.add_run()
    run._r.append(latex_to_omml_element(latex))


def fill_legend_paragraph(paragraph: Paragraph, segments: list[LegendSeg]) -> None:
    for kind, content in segments:
        if kind == "text":
            paragraph.add_run(content)
        else:
            add_inline_math_run(paragraph, content)


def add_legend_after_equation(
    equation_paragraph: Paragraph, label: str, segments: list[LegendSeg]
) -> bool:
    if next_is_legend_block(equation_paragraph, label):
        return False
    lp = insert_paragraph_after(equation_paragraph)
    fill_legend_paragraph(lp, segments)
    return True


def remove_legend_paragraphs(doc: Document) -> int:
    removed = 0
    for p in list(doc.paragraphs):
        if not is_legend_paragraph(p):
            continue
        el = p._element
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
            removed += 1
    return removed


def latex_to_omml_para(latex: str) -> str:
    mathml = latex2mathml.convert(latex)
    omml = mathml2omml.convert(mathml)
    if omml.startswith("<m:oMath"):
        inner = omml
    else:
        inner = omml
    return (
        f'<m:oMathPara xmlns:m="{MATH_NS}">'
        f'<m:oMathParaPr><m:jc m:val="center"/></m:oMathParaPr>'
        f"{inner}"
        f"</m:oMathPara>"
    )


def add_equation_after(paragraph: Paragraph, latex: str, label: str) -> bool:
    if next_is_equation_block(paragraph, label):
        return False
    fp = insert_paragraph_after(paragraph)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        fp._element.append(parse_xml(latex_to_omml_para(latex)))
    except Exception as exc:
        raise RuntimeError(f"OMML failed for {label}: {exc}") from exc
    run = fp.add_run("\t")
    run = fp.add_run(label)
    run.italic = False
    return True


def add_text_reference(paragraph: Paragraph, label: str, suffix: str) -> bool:
    if label in paragraph.text or suffix.strip() in paragraph.text:
        return False
    paragraph.add_run(suffix)
    return True


def remove_legacy_formula_paragraphs(doc: Document) -> int:
    """Удалить старые текстовые «формулы» (курсив одной строкой), если v1.4 пересобирают."""
    removed = 0
    for p in list(doc.paragraphs):
        t = p.text.strip()
        if not t:
            continue
        if "(3." in t or "(4." in t:
            if any(ch in t for ch in ("←", "Σ", "ũ", "Q_0", "ESS", "mathrm{")):
                el = p._element
                parent = el.getparent()
                if parent is not None:
                    parent.remove(el)
                    removed += 1
    return removed


def apply_formulas(doc: Document) -> list[str]:
    log: list[str] = []

    removed = remove_legacy_formula_paragraphs(doc)
    if removed:
        log.append(f"Удалено устаревших текстовых формул: {removed}")
    removed_leg = remove_legend_paragraphs(doc)
    if removed_leg:
        log.append(f"Удалено старых абзацев «где …»: {removed_leg}")

    # 1) отсылки в тексте
    for needle, _latex, label, text_suffix, _legend in FORMULA_ENTRIES:
        for p in doc.paragraphs:
            if needle in p.text:
                if add_text_reference(p, label, text_suffix):
                    log.append(f"Текст: {label} в p{paragraph_index(doc, p)}")
                break

    # 2) уравнения Word (с конца)
    tasks: list[tuple[Paragraph, str, str, list[LegendSeg]]] = []
    for needle, latex, label, _suffix, legend_segments in FORMULA_ENTRIES:
        for p in doc.paragraphs:
            if needle in p.text:
                tasks.append((p, latex, label, legend_segments))
                break
        else:
            log.append(f"Не найден абзац для {label}")

    tasks.sort(key=lambda t: paragraph_index(doc, t[0]), reverse=True)
    for p, latex, label, legend_segments in tasks:
        try:
            if add_equation_after(p, latex, label):
                log.append(f"Уравнение Word {label} после p{paragraph_index(doc, p)}")
            eq = find_equation_paragraph(p, label)
            if eq is None:
                log.append(f"Не найден блок уравнения для {label}")
                continue
            if add_legend_after_equation(eq, label, legend_segments):
                log.append(f"Расшифровка {label} (inline OMML) под уравнением")
        except Exception as exc:
            log.append(f"ОШИБКА {label}: {exc}")

    return log


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f"Не найден: {SRC}")

    shutil.copy2(SRC, DST)
    doc = Document(DST)
    log = apply_formulas(doc)
    doc.save(DST)

    lines = [
        "# Формулы Word (OMML) в `БояриновИР_Дипломная v1.4.docx`",
        "",
        f"Источник: `{SRC.name}` → `{DST.name}`",
        "",
        "Формулы вставлены через редактор уравнений Word (LaTeX → MathML → OMML).",
        "В соответствующих абзацах добавлены отсылки «см. (3.N)» / «см. (4.N)».",
        "Под каждой формулой — абзац «где …»; обозначения — inline-уравнения Word.",
        "",
        "Зависимости: `latex2mathml`, `mathml2omml`.",
        "",
        "## Журнал",
        "",
    ]
    for entry in log:
        lines.append(f"- {entry}")
    CHANGELOG.write_text("\n".join(lines), encoding="utf-8")

    print(f"Saved: {DST}")
    for entry in log:
        print(f"  - {entry}")


if __name__ == "__main__":
    main()
