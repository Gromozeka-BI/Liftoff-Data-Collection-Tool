#!/usr/bin/env python3
"""
Добавляет перекрёстные ссылки в текст черновика ВКР и подсвечивает их жёлтым
(для ручного переноса в основной документ).

Источник: БояриновИР_Дипломная v1.3.docx
Выход:   БояриновИР_ВКР_черновик.docx (или *_refs.docx, если файл открыт в Word)
"""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

HERE = Path(__file__).parent
SRC = HERE / "БояриновИР_Дипломная v1.3.docx"
DST = HERE / "БояриновИР_ВКР_черновик.docx"
CHANGELOG = HERE / "CHANGELOG_refs_highlighted.md"

REF_RE = re.compile(r"рис\.|табл\.|прил\.", re.IGNORECASE)


def norm(text: str) -> str:
    return text.replace("\xa0", " ").strip()


def highlight_run(run) -> None:
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def highlight_refs_in_paragraph(paragraph) -> bool:
    """Подсветить run'ы, в которых уже есть рис./табл./прил."""
    changed = False
    for run in paragraph.runs:
        if REF_RE.search(run.text):
            highlight_run(run)
            changed = True
    return changed


def append_suffix_highlighted(paragraph, suffix: str) -> bool:
    suffix = suffix.strip()
    if not suffix:
        return False
    full = norm(paragraph.text)
    if suffix in full:
        highlight_refs_in_paragraph(paragraph)
        return False
    run = paragraph.add_run(" " + suffix)
    highlight_run(run)
    return True


def patch_contains(paragraph, needle: str, suffix: str) -> bool:
    if needle not in paragraph.text:
        return False
    if "табл. 6" in suffix and "табл. 6" in paragraph.text:
        return False
    return append_suffix_highlighted(paragraph, suffix)


def patch_exact(paragraph, exact: str, suffix: str) -> bool:
    if norm(paragraph.text) != norm(exact):
        return False
    return append_suffix_highlighted(paragraph, suffix)


def strip_suffix_from_paragraph(paragraph, suffix: str) -> bool:
    """Удалить ошибочно вставленный фрагмент (без подсветки)."""
    suffix = suffix.strip()
    text = paragraph.text
    if suffix not in text:
        return False
    new_text = text.replace(suffix, "").replace("  ", " ").rstrip()
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)
    return True


def clean_misplaced_fig_refs(paras) -> list[str]:
    """Убрать рис. 9–10 из § 2.5 (синхронизация), если попали не в Record/Replay."""
    log: list[str] = []
    wrong_in_sync = (
        " Интерфейс Record — рис. 9; Replay — рис. 10.",
        " Интерфейс Record — рис. 9; Replay — рис. 10",
    )
    for p in paras:
        if "единую временную шкалу" in p.text:
            for suf in wrong_in_sync:
                if strip_suffix_from_paragraph(p, suf):
                    log.append("Удалена отсылка рис. 9–10 из § 2.5 (синхронизация)")
                    break
    for p in paras:
        if "основных архитектурных решений комплекса" in p.text and p.text.count("табл. 6") >= 2:
            if strip_suffix_from_paragraph(p, " Сводка временных меток (детализация) — табл. 6."):
                log.append("Убран дубликат табл. 6 в § 2.5")
    return log


def section_has_ref(paras, start: int, end: int, pattern: str) -> bool:
    for p in paras[start:end]:
        if pattern in p.text:
            return True
    return False


def fix_fig34_caption(paragraph) -> bool:
    t = paragraph.text
    if not t.startswith("Рисунок 34"):
        return False
    if "cross-drone" in t and "эксп. 11" in t:
        return False
    if "same-drone" in t and "6,49" in t:
        new = (
            "Рисунок 34 – тепловая карта mean p90 при моделировании камерного слияния "
            "(эксп. 11): условие cross-drone, сетка σ_cam × T_update "
            "(σ: 1–20 м, T: 0,2–5 с), baseline 6,49 м"
        )
        if paragraph.runs:
            paragraph.runs[0].text = new
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(new)
        highlight_refs_in_paragraph(paragraph)
        return True
    return False


# (подстрока для поиска абзаца, добавляемый суффикс)
TEXT_PATCHES: list[tuple[str, str]] = [
    # Глава 1
    (
        "требуется дополнительно рассчитывать p95.",
        "Категории HEPU/VEPU сведены в табл. 2 и иллюстрируются рис. 1.",
    ),
    (
        "должна быть предусмотрена геопривязка локальных координат",
        "Требования сведены в табл. 3; контекст системы — рис. 2.",
    ),
    (
        "либо не подходят для типового FPV-сценария.",
        "Сравнение существующих подходов — табл. 1.",
    ),
    # Глава 2
    (
        "Верхнеуровневая структура комплекса включает четыре логических блока",
        "Состав блоков — табл. 4; диаграмма контейнеров — рис. 3.",
    ),
    (
        "что позволило сопоставлять результаты разработки между виртуальной и физической средой.",
        "Аппаратура испытаний — рис. 4.",
    ),
    (
        "передаёт его на ПК через COM-порт.",
        "Состав аппаратной части — табл. 5; схема RC-цепочки — рис. 5–6.",
    ),
    (
        "для сопоставления RC-каналов с видеопотоком и телеметрией.",
        "Временные метки потоков — табл. 6; схема сбора RC — рис. 7.",
    ),
    (
        "камерное наблюдение связано с моментом времени.",
        "Схема видеотракта — рис. 8.",
    ),
    (
        "основных архитектурных решений комплекса.",
        "Сводка временных меток — табл. 6.",
    ),  # не добавляется, если табл. 6 уже есть в абзаце
    # § 2.7 Record / Replay — рис. 9–10 (см. patch_record_replay_refs)
    (
        "«CameraObservation». Благодаря этому основной контур может работать без камеры",
        "",  # уже есть ссылки в v1.3 — только подсветка
    ),
    (
        "В нём выделены три ключевые сущности",
        "Сущности алгоритма — табл. 9.",
    ),
    # § 3.2 — рис. 11 (если ещё нет в абзаце про .npz)
    (
        "может ухудшать переносимость.",
        "Признаки локализации — табл. 10.",
    ),
    # Глава 3 — эксперименты (общее)
    (
        "Внутренний критерий работоспособности был задан как «p90 < 15 м».",
        "Перечень экспериментов — табл. 13.",
    ),
    (
        "определена структура наблюдения.",
        "Сводка первого блока — табл. 14.",
    ),
    (
        "перестраивать референс и перенастраивать локализатор.",
        "Сводка второго блока — табл. 23.",
    ),
    (
        "Блок экспериментов 11–12 посвящён использованию видеопотока",
        "Сводка четвёртого блока — табл. 36.",
    ),
    # Глава 4
    (
        "«source_system» и «source_component».",
        "Интеграционные компоненты — табл. 41; интерфейс Nebosvod Connect — рис. 44.",
    ),
    (
        "передаёт информацию в систему «Небосвод».",
        "",  # рис. 39–40 уже в тексте
    ),
    (
        "требуется отдельная оценка p95.",
        "Критерии HEPU/VEPU — табл. 43.",
    ),
    (
        "нужно рассчитывать p95 на репрезентативной выборке.",
        "Направления развития — табл. 44.",
    ),
    (
        "Десятый эксперимент проверял, сохраняется ли работоспособность RC/PF-контура",
        "Подробности — табл. 34–35 (см. также выводы эксперимента).",
    ),
]

# (табл. условий, табл. результатов, рисунки) для экспериментов 1–12
EXP_TABLES_FIGS: list[tuple[int, int, str]] = [
    (15, 16, "14"),
    (17, 18, "15"),
    (19, 20, "16–17"),
    (21, 22, "18–19"),
    (24, 25, "20–21"),
    (26, 27, "22–23"),
    (28, 29, "24"),
    (30, 31, "25–26"),
    (32, 33, "27–28"),
    (34, 35, "29–32"),
    (37, 38, "33–35"),
    (39, 40, "36–38"),
]

APP_HEADING_SUFFIX = {
    "ПРИЛОЖЕНИЕ Е. Блок схема алгоритма Particle Filter": " (§ 3.4, рис. 45, табл. 11)",
    "ПРИЛОЖЕНИЕ Ж. Блок схема алгоритма позиционирования по видеопотоку": " (§ 3.5, рис. 46, табл. 12)",
    "ПРИЛОЖЕНИЕ Н. Диаграмма компонентов DCT Desktop (C4, Component)": " (§ 2.6, рис. 49)",
    "ПРИЛОЖЕНИЕ З. Фрагмент inject_position_observation": " (§ 3.6)",
    "ПРИЛОЖЕНИЕ И. Пайплайн формирования camera_observations.jsonl": " (§ 3.5, рис. 46)",
    "ПРИЛОЖЕНИЕ А. Структура выходной сессии DCT": " (§ 2.6, табл. 8)",
}


def collect_exp_paragraph_indices(paras) -> tuple[list[int], list[int]]:
    """Индексы абзацев сразу после «Описание» и «Выводы» в каждом эксперименте."""
    desc_idxs: list[int] = []
    vyv_idxs: list[int] = []
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t == "Описание" and i + 1 < len(paras):
            desc_idxs.append(i + 1)
        if t == "Выводы" and i + 1 < len(paras):
            vyv_idxs.append(i + 1)
    return desc_idxs, vyv_idxs


def patch_experiment_table_refs(paras) -> list[str]:
    """Отсылки на табл. 15–40 в тексте экспериментов (Описание + Выводы)."""
    log: list[str] = []
    desc_idxs, vyv_idxs = collect_exp_paragraph_indices(paras)

    for n, (tab_cond, tab_res, fig) in enumerate(EXP_TABLES_FIGS, start=1):
        if n > len(desc_idxs) or n > len(vyv_idxs):
            break
        desc_p = paras[desc_idxs[n - 1]]
        vyv_p = paras[vyv_idxs[n - 1]]

        desc_suffix = f" Условия и методика — табл. {tab_cond}."
        if f"табл. {tab_cond}" not in desc_p.text and append_suffix_highlighted(
            desc_p, desc_suffix
        ):
            log.append(f"Эксп. {n} Описание: табл. {tab_cond}")

        if f"табл. {tab_res}" not in vyv_p.text:
            vyv_suffix = (
                f" Таблицы: условия — табл. {tab_cond}, результаты — табл. {tab_res}; "
                f"иллюстрации — рис. {fig}."
            )
            if append_suffix_highlighted(vyv_p, vyv_suffix):
                log.append(f"Эксп. {n} Выводы: табл. {tab_cond}–{tab_res}, рис. {fig}")
        elif f"рис. {fig.split('–')[0].strip()}" not in vyv_p.text:
            if append_suffix_highlighted(vyv_p, f" Иллюстрации — рис. {fig}."):
                log.append(f"Эксп. {n} Выводы: рис. {fig}")

    return log


def patch_record_replay_refs(paras) -> list[str]:
    """Рис. 9–10 — в § 2.7, один раз на режим."""
    log: list[str] = []
    record_markers = (
        ("согласованном формате", " (рис. 9)."),
        (
            "качество оценивалось по дискретным событиям прохождения ворот.",
            " Интерфейс режима Record — рис. 9.",
        ),
    )
    if not section_has_ref(paras, 150, 158, "рис. 9"):
        for needle, suffix in record_markers:
            for p in paras:
                if needle in p.text and append_suffix_highlighted(p, suffix):
                    log.append(f"§ 2.7 Record: {suffix.strip()}")
                    break

    if not section_has_ref(paras, 158, 165, "рис. 10"):
        for p in paras:
            if "подключать их к локализатору." in p.text and append_suffix_highlighted(
                p, " Интерфейс режима Replay — рис. 10."
            ):
                log.append("§ 2.7 Replay: рис. 10")
                break
    return log


def patch_paragraphs(doc: Document) -> list[str]:
    log: list[str] = []
    paras = doc.paragraphs

    log.extend(clean_misplaced_fig_refs(paras))
    log.extend(patch_record_replay_refs(paras))
    log.extend(patch_experiment_table_refs(paras))

    for p in paras:
        if fix_fig34_caption(p):
            log.append("Рис. 34: подпись cross-drone (эксп. 11)")

    if not any("рис. 11" in p.text for p in paras[178:188]):
        for p in paras:
            if "между текущим наблюдением и узлами референса." in p.text:
                if append_suffix_highlighted(p, " Схема представления трассы — рис. 11."):
                    log.append("§ 3.2: рис. 11 у референсной траектории")
                break

    for needle, suffix in TEXT_PATCHES:
        if not suffix:
            for p in paras:
                if needle in p.text and highlight_refs_in_paragraph(p):
                    log.append(f"Подсветка (уже в тексте): …{needle[:40]}…")
                    break
            continue
        for p in paras:
            if patch_contains(p, needle, suffix):
                log.append(f"Добавлено: …{needle[:36]}… -> {suffix[:50]}")
                break

    # § 3.4–3.6, § 4 — подсветить уже встроенные ссылки v1.3
    for marker in (
        "Параметры алгоритма — табл. 11",
        "Блок-схема конвейера — рис. 46",
        "Поля наблюдения — табл. 12",
        "Фрагмент реализации — прил",
        "Схема потока данных — рис. 40",
        "Параметры преобразования — табл. 42",
        "Последовательность этапов — рис. 43",
        "Подсистемы DCT — табл",
    ):
        for p in paras:
            if marker in p.text:
                if highlight_refs_in_paragraph(p):
                    log.append(f"Подсветка v1.3: {marker}")

    for p in paras:
        t = p.text.strip()
        if t in APP_HEADING_SUFFIX:
            suf = APP_HEADING_SUFFIX[t]
            if suf not in p.text:
                run = p.add_run(suf)
                highlight_run(run)
                log.append(f"Прил. заголовок: {t[:40]}…")

    return log


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f"Не найден: {SRC}")

    tmp = DST.with_suffix(".tmp.docx")
    shutil.copy2(SRC, tmp)
    doc = Document(tmp)
    log = patch_paragraphs(doc)

    out = DST
    try:
        doc.save(out)
        tmp.unlink(missing_ok=True)
    except PermissionError:
        out = DST.with_name(DST.stem + "_refs.docx")
        doc.save(out)
        tmp.unlink(missing_ok=True)
        print(
            f"WARNING: {DST.name} открыт в Word — сохранено в {out.name}. "
            "Закройте файл и переименуйте или запустите снова."
        )

    lines = [
        "# Ссылки с подсветкой в черновике",
        "",
        f"Источник: `{SRC.name}` → `{out.name}`",
        "",
        "Добавленные и уже существовавшие фрагменты с «рис.» / «табл.» / «прил.» "
        "подсвечены **жёлтым** (Word: цвет выделения).",
        "",
        "**v2:** рис. 9–10 — в § 2.7 (Record/Replay); рис. 11 — в § 3.2 (референсная траектория).",
        "",
        "**v3:** в каждом эксперименте 1–12 — табл. условий/методики и табл. результатов "
        "в «Описание» и «Выводы» + рис.",
        "",
        "## Журнал",
        "",
    ]
    for entry in log:
        lines.append(f"- {entry}")
    if not log:
        lines.append("- (изменений не зафиксировано)")
    CHANGELOG.write_text("\n".join(lines), encoding="utf-8")

    print(f"Saved: {out}")
    print(f"Entries: {len(log)}")
    for entry in log:
        print(f"  - {entry}")


if __name__ == "__main__":
    main()
