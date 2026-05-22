#!/usr/bin/env python3
"""
Обновление БояриновИР_ВКР_черновик.docx на основе v1.2:
- копия содержимого v1.2;
- перекрёстные ссылки встраиваются в текст абзацев (не отдельными строками у таблиц).
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

HERE = Path(__file__).parent
SRC = HERE / "БояриновИР_Дипломная v1.2.docx"
DST = HERE / "БояриновИР_ВКР_черновик.docx"
CHANGELOG = HERE / "CHANGELOG_v13_черновик.md"


def set_paragraph_text(paragraph, text: str) -> None:
    """Надёжная замена текста абзаца (в т.ч. при нескольких w:r)."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def append_suffix(paragraph, suffix: str) -> bool:
    """Добавить фразу в конец абзаца, если её ещё нет."""
    suffix = suffix.strip()
    if not suffix or suffix in paragraph.text:
        return False
    set_paragraph_text(paragraph, paragraph.text.rstrip() + " " + suffix)
    return True


def replace_exact(paragraph, old: str, new: str) -> bool:
    if paragraph.text != old:
        return False
    set_paragraph_text(paragraph, new)
    return True


def replace_contains(paragraph, old: str, new: str) -> bool:
    if old not in paragraph.text:
        return False
    set_paragraph_text(paragraph, paragraph.text.replace(old, new, 1))
    return True


def patch_paragraphs(doc: Document) -> list[str]:
    log: list[str] = []
    paras = doc.paragraphs

    # --- Рис. 33–34 (эксп. 11) ---
    for p in paras:
        t = p.text
        if t.startswith("Рисунок 34") and "same-drone" in t and "6,49" in t:
            set_paragraph_text(
                p,
                "Рисунок 34 – тепловая карта mean p90 при моделировании камерного слияния "
                "(эксп. 11): условие cross-drone, сетка σ_cam × T_update "
                "(σ: 1–20 м, T: 0,2–5 с), baseline 6,49 м",
            )
            log.append("Рис. 34: подпись cross-drone")
        if t.startswith("Рисунок 33") and "эксп. 11" not in t:
            set_paragraph_text(
                p,
                "Рисунок 33 – тепловая карта mean p90 при моделировании камерного слияния "
                "(эксп. 11): условие same-drone, сетка σ_cam × T_update "
                "(σ: 1–20 м, T: 0,2–5 с), baseline 12,36 м",
            )
            log.append("Рис. 33: уточнена подпись (эксп. 11)")

    # --- Глава 1: ссылки в тексте перед таблицами ---
    ch1_refs = [
        (
            "Поэтому результаты, полученные в исследовательской части по метрике p90, "
            "должны рассматриваться как внутренняя экспериментальная оценка, а для строгого "
            "сопоставления с HEPU требуется дополнительно рассчитывать p95.",
            "Категории HEPU/VEPU сведены в табл. 2 и иллюстрируются рис. 1.",
            "§ 1 (HEPU): табл. 2, рис. 1 в тексте",
        ),
        (
            "Для интеграции с внешней системой мониторинга должна быть предусмотрена "
            "геопривязка локальных координат и передача результата в формате, пригодном "
            "для дальнейшей обработки.",
            "Требования сведены в табл. 3; контекст системы — рис. 2.",
            "§ 1 (требования): табл. 3, рис. 2 в тексте",
        ),
        (
            "Верхнеуровневая структура комплекса включает четыре логических блока: блок "
            "получения RC-данных, блок получения видеопотока, программный блок DCT и блок "
            "интеграции с внешней системой мониторинга. Внутри DCT выделены режимы Record "
            "и Replay. Первый отвечает за запись сессии, второй - за повторный анализ, "
            "отладку локализации и воспроизведение экспериментов.",
            "Состав блоков — табл. 4; диаграмма контейнеров — рис. 3.",
            "§ 2.1: табл. 4, рис. 3 в тексте",
        ),
    ]
    for anchor, suffix, msg in ch1_refs:
        for p in paras:
            if p.text == anchor and append_suffix(p, suffix):
                log.append(msg)
                break

    # --- § 2.6 ---
    ref_26_anchor = (
        "Отдельное место занимает пакет камерной локализации. Он не является жёсткой "
        "зависимостью основного RC-локализатора: связь выполняется через готовое наблюдение "
        "«CameraObservation». Благодаря этому основной контур может работать без камеры, "
        "а камерный модуль подключается как дополнительный источник абсолютных координат."
    )
    ref_26_suffix = (
        "Подсистемы DCT — табл. 7; детализация модулей — рис. 49 (прил. Н); "
        "структура сессии — табл. 8 и прил. А."
    )
    for p in paras:
        if p.text == ref_26_anchor and append_suffix(p, ref_26_suffix):
            log.append("§ 2.6: табл. 7, рис. 49, табл. 8, прил. А в тексте")
            break

    # --- § 3.4 PF ---
    old_pf_open = (
        "Основным рабочим алгоритмом выбран Particle Filter (Приложение Е). "
        "В реализации DCT каждая частица имеет состояние «(s, v)», где «s» - позиция "
        "вдоль дуги трассы в метрах, а «v» - скорость движения вдоль референсной траектории. "
        "Такое состояние достаточно компактно, но сохраняет важную динамическую информацию: "
        "фильтр не просто сопоставляет текущий кадр с референсом, а поддерживает гипотезы "
        "о положении и скорости."
    )
    new_pf_open = (
        "Основным рабочим алгоритмом выбран Particle Filter. "
        "В реализации DCT каждая частица имеет состояние «(s, v)», где «s» - позиция "
        "вдоль дуги трассы в метрах, а «v» - скорость движения вдоль референсной траектории. "
        "Такое состояние достаточно компактно, но сохраняет важную динамическую информацию: "
        "фильтр не просто сопоставляет текущий кадр с референсом, а поддерживает гипотезы "
        "о положении и скорости."
    )
    pf_close_anchor = (
        "После обновления весов вычисляется эффективное число частиц ESS. Если "
        "ESS становится меньше заданной доли от общего числа частиц, выполняется "
        "систематический ресемплинг. Для сохранения способности к восстановлению "
        "после ошибочной локализации добавляется roughening по координате «s», "
        "а в обычном RC-обновлении часть частиц может случайно переинициализироваться "
        "по всей длине трассы. Оценка положения вычисляется как circular mean по "
        "частицам, что корректно обрабатывает границу старт/финиш."
    )
    pf_close_suffix = (
        "Параметры алгоритма — табл. 11; полный цикл одного шага — рис. 45 (прил. Е)."
    )

    for p in paras:
        if replace_exact(p, old_pf_open, new_pf_open):
            log.append("§ 3.4: убрана ранняя отсылка «(Приложение Е)»")
        if p.text == pf_close_anchor and append_suffix(p, pf_close_suffix):
            log.append("§ 3.4: табл. 11 и рис. 45 (прил. Е) в тексте (перед табл. 11)")
            break

    # --- § 3.5 Camera ---
    old_cam_end = (
        "После получения детекций выполняется сопоставление с моделью ворот трассы "
        "и решение задачи PnP. В коде эта часть реализована через «YoloGateDetection», "
        "«YoloAdapterConfig» и «CoarseRefineLocalizer». На выходе формируется "
        "«CameraObservation»: позиция «xyz_obs» в локальной системе координат трассы, "
        "оценка неопределённости «sigma_cam», «confidence», «gate_id», ошибка репроекции, "
        "номер кадра, «bbox» и «keypoints» для последующей визуализации в Replay (Приложение Ж)."
    )
    new_cam_end = (
        "После получения детекций выполняется сопоставление с моделью ворот трассы "
        "и решение задачи PnP. В коде эта часть реализована через «YoloGateDetection», "
        "«YoloAdapterConfig» и «CoarseRefineLocalizer». На выходе формируется "
        "«CameraObservation»: позиция «xyz_obs» в локальной системе координат трассы, "
        "оценка неопределённости «sigma_cam», «confidence», «gate_id», ошибка репроекции, "
        "номер кадра, «bbox» и «keypoints» для последующей визуализации в Replay. "
        "Поля наблюдения — табл. 12; блок-схема offline-конвейера (YOLO → PnP → jsonl) — "
        "рис. 46 (прил. Ж); примеры ворот и разметки — рис. 12–13."
    )
    cam_open_anchor = (
        "Камерный модуль формирует абсолютные наблюдения положения по видеопотоку "
        "курсовой камеры. На текущем этапе используется offline-цепочка: из папки "
        "записанной сессии читается «video.mp4», временные метки кадров берутся "
        "из «video_timestamps.parquet», после чего на выбранных кадрах запускается "
        "YOLO-модель для обнаружения ворот и четырёх внутренних углов."
    )
    cam_open_suffix = "Блок-схема конвейера — рис. 46 (прил. Ж)."

    cam_pnp_done = False
    for p in paras:
        if (
            not cam_pnp_done
            and "CoarseRefineLocalizer" in p.text
            and "Replay" in p.text
            and "CameraObservation" in p.text
        ):
            if "Приложение Ж" in p.text:
                set_paragraph_text(p, new_cam_end)
            elif "рис. 46" not in p.text:
                append_suffix(
                    p,
                    "Поля наблюдения — табл. 12; блок-схема offline-конвейера — "
                    "рис. 46 (прил. Ж); примеры — рис. 12–13.",
                )
            log.append("§ 3.5: табл. 12, рис. 46, рис. 12–13 в тексте (абзац про PnP)")
            cam_pnp_done = True
        if p.text == cam_open_anchor and append_suffix(p, cam_open_suffix):
            log.append("§ 3.5: рис. 46 в тексте (offline-цепочка)")

    # --- § 3.6 ---
    inject_old = (
        "Слияние RC-локализации и камерных наблюдений реализовано как дополнительное "
        "байесовское обновление весов частиц. Метод «inject_position_observation» принимает "
        "абсолютное XYZ-наблюдение и «sigma_cam». Для каждой частицы вычисляется её 3D-позиция "
        "на референсе, после чего расстояние до камерного наблюдения преобразуется в правдоподобие "
        "«exp(-0.5 * d^2 / sigma_cam^2)»."
    )
    inject_new = inject_old + " Фрагмент реализации — прил. З."
    for p in paras:
        if p.text == inject_old:
            set_paragraph_text(p, inject_new)
            log.append("§ 3.6: прил. З в тексте")
            break

    # --- § 4.1 ---
    for p in paras:
        if (
            "интеграция реализуется через отдельное приложение Nebosvod Connect" in p.text
            and "рис. 39" not in p.text
        ):
            append_suffix(p, "Схема потока данных — рис. 40; пример экрана — рис. 39.")
            log.append("§ 4.1: рис. 39–40 в тексте")
            break

    # --- § 4.2 ---
    geo_anchor = (
        "В DCT эта задача решается через ручную геопривязку по трём точкам: "
        "«origin», «x» и «z». Точка «origin» задаёт географическое положение начала "
        "локальной области. Точка «x» задаёт направление и масштаб локальной оси X, "
        "а точка «z» – направление и масштаб локальной оси Z. Пользователь "
        "самостоятельно вычисляет координаты этих точек для конкретного полигона или трассы."
    )
    geo_suffix = (
        "Параметры преобразования — табл. 42; схема точек — рис. 41; пример экрана — рис. 42."
    )
    for p in paras:
        if p.text == geo_anchor and append_suffix(p, geo_suffix):
            log.append("§ 4.2: табл. 42, рис. 41–42 в тексте геопривязки")
            break

    # --- § 4.3 ---
    transform_anchor = (
        "Преобразование реализовано в модуле «dct.mavlink_geo». Сначала географические "
        "точки переводятся в локальную касательную систему ENU относительно «origin». "
        "Затем по точкам «x» и «z» вычисляются базисные векторы, соответствующие единице "
        "локальных координат трассы. После этого произвольная точка «[x, y, z]» из DCT "
        "переводится в ENU и обратно в широту, долготу и высоту."
    )
    transform_suffix = "Последовательность этапов — рис. 43."
    for p in paras:
        if p.text == transform_anchor and append_suffix(p, transform_suffix):
            log.append("§ 4.3: рис. 43 в тексте преобразования")
            break

    # --- Приложения: дополнение заголовков (не отдельный абзац) ---
    app_heading_suffix = {
        "ПРИЛОЖЕНИЕ Е. Блок схема алгоритма Particle Filter": (
            " (§ 3.4, рис. 45, табл. 11)"
        ),
        "ПРИЛОЖЕНИЕ Ж. Блок схема алгоритма позиционирования по видеопотоку": (
            " (§ 3.5, рис. 46, табл. 12)"
        ),
        "ПРИЛОЖЕНИЕ Н. Диаграмма компонентов DCT Desktop (C4, Component)": (
            " (§ 2.6, рис. 49)"
        ),
        "ПРИЛОЖЕНИЕ З. Фрагмент inject_position_observation": " (§ 3.6)",
        "ПРИЛОЖЕНИЕ И. Пайплайн формирования camera_observations.jsonl": (
            " (§ 3.5, рис. 46)"
        ),
    }
    for p in paras:
        t = p.text.strip()
        if t in app_heading_suffix and app_heading_suffix[t] not in p.text:
            set_paragraph_text(p, t + app_heading_suffix[t])
            log.append(f"Прил.: ссылка в заголовке «{t[:36]}…»")

    return log


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f"Не найден исходный файл: {SRC}")

    out = DST
    tmp = DST.with_suffix(".tmp.docx")
    shutil.copy2(SRC, tmp)
    doc = Document(tmp)
    log = patch_paragraphs(doc)
    try:
        doc.save(out)
        tmp.unlink(missing_ok=True)
    except PermissionError:
        out = DST.with_name(DST.stem + "_v13.docx")
        doc.save(out)
        print(
            f"WARNING: {DST.name} открыт в Word — сохранено в {out.name}. "
            "Закройте файл и переименуйте или запустите скрипт снова."
        )
        tmp.unlink(missing_ok=True)

    lines = [
        "# Изменения в `БояриновИР_ВКР_черновик.docx` относительно v1.2",
        "",
        f"Источник: `{SRC.name}` → `{out.name}`",
        "",
        "**Принцип:** все перекрёстные ссылки встроены в текст существующих абзацев, "
        "без отдельных строк между таблицей и её подписью.",
        "",
        "## Затронутые разделы",
        "",
    ]
    sections = {
        "§ 3.4": "Глава 3 — Particle Filter",
        "§ 3.5": "Глава 3 — камерные наблюдения",
        "§ 3.6": "Глава 3 — слияние RC и камеры",
        "§ 2.6": "Глава 2 — структура DCT",
        "§ 4.1": "Глава 4 — интеграция",
        "§ 4.2": "Глава 4 — геопривязка",
        "§ 4.3": "Глава 4 — преобразование координат",
        "§ 1": "Глава 1",
        "§ 2.1": "Глава 2 — архитектура",
        "Рис. 33": "Глава 3 — эксп. 11",
        "Рис. 34": "Глава 3 — эксп. 11",
        "Прил.": "Приложения",
    }
    touched: set[str] = set()
    for entry in log:
        for key, title in sections.items():
            if key in entry:
                touched.add(title)
    for title in sorted(touched):
        lines.append(f"- **{title}**")
    lines.extend(["", "## Детальный журнал правок", ""])
    for entry in log:
        lines.append(f"- {entry}")
    lines.extend(
        [
            "",
            "## Что сделать вручную в Word (v1.3)",
            "",
            "1. Обновить поле **Содержание** (ПКМ → Обновить поле).",
            "2. При необходимости оформить гиперссылки на «табл. N» / «рис. N».",
            "",
        ]
    )
    CHANGELOG.write_text("\n".join(lines), encoding="utf-8")
    print(f"Saved: {out}")
    print(f"Changelog: {CHANGELOG}")
    for entry in log:
        print(f"  - {entry}")


if __name__ == "__main__":
    main()
