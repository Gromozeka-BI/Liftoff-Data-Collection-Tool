#!/usr/bin/env python3
"""
Копия БояриновИР_Дипломная v1.5.docx → v1.6:
обновлённый «Перечень сокращений и обозначений» (группы + недостающие термины).
v1.5 не изменяется.
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

HERE = Path(__file__).parent
SRC = HERE / "БояриновИР_Дипломная v1.5.docx"
DST = HERE / "БояриновИР_Дипломная v1.6.docx"
CHANGELOG = HERE / "CHANGELOG_v16_перечень.md"

# (текст абзаца, жирный заголовок группы?)
AbbrevLine = tuple[str, bool]

ABBREV_SECTION: list[AbbrevLine] = [
    (
        "Ниже приведены сокращения и обозначения, используемые в работе. "
        "Сокращения сгруппированы по тематике; математические обозначения "
        "согласованы с формулами глав 3–4.",
        False,
    ),
    ("Сокращения", True),
    ("Общие термины и предметная область", True),
    ("БВС — беспилотное воздушное судно.", False),
    (
        "FPV (First Person View) — управление и наблюдение «от первого лица»; "
        "в работе — FPV-дрон исследуемого класса.",
        False,
    ),
    ("GPS — Global Positioning System, глобальная спутниковая навигация.", False),
    ("ГЛОНАСС — глобальная навигационная спутниковая система.", False),
    ("SLAM — simultaneous localization and mapping.", False),
    ("VIO — visual-inertial odometry, визуально-инерциальная одометрия.", False),
    ("INS — inertial navigation system, инерциальная навигационная система.", False),
    ("Каналы радиоуправления (оси стиков)", True),
    (
        "Pitch — тангаж; Roll — крен; Yaw — рысканье; Throttle — газ "
        "(среднее управление тягой). В тексте — RC-каналы [throttle, yaw, pitch, roll].",
        False,
    ),
    ("Радиоуправление, приём данных и полётный контроллер", True),
    ("RC — Radio Control, радиоуправление; в работе — RC-каналы, RC-признаки.", False),
    (
        "ELRS (ExpressLRS) — протокол радиоуправления с открытым исходным кодом "
        "для FPV-дронов.",
        False,
    ),
    (
        "FLRC (Frequency-Locked Rate Control) — режим управления скоростью "
        "с частотной блокировкой в системе ELRS.",
        False,
    ),
    (
        "EdgeTX — прошивка/экосистема аппаратуры радиоуправления на стороне пилота.",
        False,
    ),
    (
        "Betaflight — открытое ПО полётного контроллера (профили rate, expo, суперрейт).",
        False,
    ),
    (
        "ESP32 — микроконтроллер приёмника RC-данных (передача на ПК по COM-порту).",
        False,
    ),
    ("COM — последовательный порт ПК для приёма потока RC.", False),
    ("Цифровой видеотракт FPV", True),
    (
        "HDZero — экосистема цифрового FPV-видеолинка (передатчик VTX и приёмник VRX).",
        False,
    ),
    ("VTX — video transmitter, передатчик видеосигнала с курсовой камеры.", False),
    ("VRX — video receiver, приёмник видеосигнала (HDZero EventVRX).", False),
    ("Программный комплекс и интерфейсы", True),
    ("DCT — Data Collection Toolkit, разработанный программный комплекс.", False),
    ("Record — режим DCT записи сессии полёта.", False),
    ("Replay — режим DCT воспроизведения и анализа записанной сессии.", False),
    ("CLI — Command Line Interface, интерфейс командной строки.", False),
    ("API — Application Programming Interface, программный интерфейс.", False),
    ("C4 — нотация архитектурных диаграмм (Context, Container, Component).", False),
    ("Алгоритмы локализации, зрения и сглаживания", True),
    ("PF — Particle Filter, фильтр частиц (основной алгоритм локализации).", False),
    ("ESS — effective sample size, эффективное число частиц.", False),
    (
        "KF — Kalman filter, фильтр Калмана (второй контур сглаживания KFLayer2).",
        False,
    ),
    (
        "CamKF — экспериментальный контур камерного сглаживания в Replay.",
        False,
    ),
    ("YOLO — семейство нейросетевых моделей детектирования объектов.", False),
    (
        "PnP — Perspective-n-Point, восстановление положения камеры "
        "по 2D–3D соответствиям.",
        False,
    ),
    (
        "RC+Rate — режим признаков: RC-стики с преобразованием в угловые скорости "
        "по профилю Betaflight.",
        False,
    ),
    ("Координаты, протоколы и внешняя система", True),
    ("ENU — East–North–Up, локальная геодезическая система координат.", False),
    ("WGS-84 — World Geodetic System 1984, географические координаты.", False),
    ("MAVLink — протокол обмена сообщениями для беспилотных систем.", False),
    ("UDP — User Datagram Protocol (транспорт MAVLink в NebosvodConnect).", False),
    (
        "NebosvodConnect — приложение интеграции DCT с системой мониторинга «Небосвод».",
        False,
    ),
    ("Метрики, категории точности и форматы данных", True),
    ("HEPU — категория горизонтальной погрешности навигационной точности.", False),
    ("VEPU — категория вертикальной погрешности навигационной точности.", False),
    ("СППИ — система представления полётной информации.", False),
    ("p90_err_m — 90-й процентиль ошибки локализации, м.", False),
    ("p95_err_m — 95-й процентиль ошибки локализации, м (сопоставление с HEPU).", False),
    ("GT (ground truth) — эталонная (истинная) траектория в экспериментах.", False),
    ("JSONL — JSON Lines (файл camera_observations.jsonl).", False),
    ("NPZ — сжатый архив NumPy (референс трассы, .npz).", False),
    ("Обозначения", True),
    (
        "s — дуговая координата вдоль референсной трассы, м; v — скорость вдоль трассы, м/с; "
        "L — длина референсного круга, м.",
        False,
    ),
    (
        "s_i, v_i, w_i — положение, скорость и вес i-й частицы PF; Δt — шаг по времени, с; "
        "η_v, η_s — возмущения скорости и положения (process_noise_v, process_noise_s).",
        False,
    ),
    (
        "u_k, ũ_k — k-й признак наблюдения и его z-score; μ_k, σ_k — среднее и СКО признака "
        "по эталонному кругу; K — число признаков (каналов).",
        False,
    ),
    (
        "d_i — невязка RC-признаков частицы i; σ_obs — параметр правдоподобия (obs_sigma); "
        "ESS — эффективное число частиц; θ_i, s_est — угол на окружности и оценка положения.",
        False,
    ),
    (
        "z_obs, σ_cam — камерное наблюдение и неопределённость (xyz_obs, sigma_cam); "
        "p(s_i) — 3D-позиция частицы на референсе.",
        False,
    ),
    (
        "e_t — модуль ошибки локализации в момент t; p_gt, p_est — истинная и оценочная "
        "позиции; Q_0,9, Q_0,95 — квантили выборки ошибок.",
        False,
    ),
    (
        "East, North, Up — компоненты ENU; φ_0, Δλ, Δφ, Δh — широта опорной точки и "
        "приращения долготы, широты, высоты; x_loc, z_loc — координаты в локальной СК трассы.",
        False,
    ),
    (
        "T_k — порог k-й категории HEPU; P(·) — вероятность попадания ошибки в допуск "
        "(оценка по p95_err_m).",
        False,
    ),
]


def find_section_bounds(doc: Document) -> tuple[Paragraph, Paragraph]:
    title: Paragraph | None = None
    after: Paragraph | None = None
    for p in doc.paragraphs:
        t = p.text.strip().upper()
        if title is None and "ПЕРЕЧЕНЬ СОКРАЩЕНИЙ" in t:
            title = p
            continue
        if title is not None and t == "ВВЕДЕНИЕ":
            after = p
            break
    if title is None or after is None:
        raise RuntimeError("Не найден блок «Перечень …» / «ВВЕДЕНИЕ»")
    return title, after


def clear_between(title: Paragraph, stop: Paragraph) -> int:
    removed = 0
    el = title._element.getnext()
    while el is not None and el is not stop._element:
        nxt = el.getnext()
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
            removed += 1
        el = nxt
    return removed


def insert_before(anchor: Paragraph, text: str, *, bold: bool = False) -> Paragraph:
    new_el = OxmlElement("w:p")
    anchor._element.addprevious(new_el)
    para = Paragraph(new_el, anchor._parent)
    run = para.add_run(text)
    run.bold = bold
    return para


def fill_abbrev_section(doc: Document) -> int:
    title, intro = find_section_bounds(doc)
    removed = clear_between(title, intro)
    anchor = intro
    for text, is_bold in ABBREV_SECTION:
        insert_before(anchor, text, bold=is_bold)
    return removed


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f"Не найден: {SRC}")

    shutil.copy2(SRC, DST)
    doc = Document(DST)
    removed = fill_abbrev_section(doc)
    doc.save(DST)

    lines = [
        "# Перечень сокращений — `БояриновИР_Дипломная v1.6.docx`",
        "",
        f"Источник: `{SRC.name}` → `{DST.name}` (v1.5 не изменялся).",
        "",
        "## Что было в v1.5",
        "",
        "- 21 абзац без групп: оси Pitch/Roll/Yaw/Throttle, FPV, ELRS, FLRC, Betaflight, CLI, API, ESP32, затем БВС…СППИ.",
        "- Не было: GPS/ГЛОНАСС, EdgeTX, HDZero/VTX/VRX, Record/Replay, ENU/WGS-84, ESS/KF/CamKF, NebosvodConnect, метрик p90/p95, форматов JSONL/NPZ, блока **обозначений**.",
        "",
        "## Что сделано в v1.6",
        "",
        f"- Удалено старых абзацев перечня: {removed}.",
        f"- Добавлено новых абзацев: {len(ABBREV_SECTION)}.",
        "- Сокращения сгруппированы по 10 тематическим подзаголовкам.",
        "- Добавлен раздел **Обозначения** (символы из формул (3.1)–(4.2)).",
        "- Pitch/Roll/Yaw/Throttle объединены в один абзац про RC-каналы.",
        "",
        "## Группы сокращений",
        "",
        "1. Общие термины (БВС, FPV, GPS, ГЛОНАСС, SLAM, VIO, INS)",
        "2. Каналы RC",
        "3. Радиоуправление и ПК (RC, ELRS, FLRC, EdgeTX, Betaflight, ESP32, COM)",
        "4. Видеотракт (HDZero, VTX, VRX)",
        "5. DCT и интерфейсы (DCT, Record, Replay, CLI, API, C4)",
        "6. Алгоритмы (PF, ESS, KF, CamKF, YOLO, PnP, RC+Rate)",
        "7. Координаты и интеграция (ENU, WGS-84, MAVLink, UDP, NebosvodConnect)",
        "8. Метрики и данные (HEPU, VEPU, СППИ, p90/p95, GT, JSONL, NPZ)",
        "9. Обозначения (s, v, L, частицы, z-score, PF, камера, метрики, ENU, HEPU)",
        "",
    ]
    CHANGELOG.write_text("\n".join(lines), encoding="utf-8")
    print(f"Saved: {DST}")
    print(f"Removed old lines: {removed}, added: {len(ABBREV_SECTION)}")


if __name__ == "__main__":
    main()
