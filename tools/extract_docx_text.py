from __future__ import annotations

import sys
from pathlib import Path

import docx


def extract(path: Path) -> str:
    d = docx.Document(str(path))
    out: list[str] = []

    for para in d.paragraphs:
        t = (para.text or "").rstrip()
        if t:
            out.append(t)

    for ti, table in enumerate(d.tables):
        out.append("")
        out.append(f"[TABLE {ti + 1}]")
        for row in table.rows:
            cells = [(c.text or "").replace("\n", " ").strip() for c in row.cells]
            out.append("\t".join(cells))

    return "\n".join(out).rstrip() + "\n"


def main(argv: list[str]) -> int:
    sys.stdout.reconfigure(encoding="utf-8")
    if len(argv) != 2:
        print("Usage: python tools/extract_docx_text.py <path.docx>", file=sys.stderr)
        return 2

    p = Path(argv[1])
    if not p.exists():
        print(f"File not found: {p}", file=sys.stderr)
        return 1

    print("=" * 120)
    print(f"FILE: {p.name}")
    print("=" * 120)
    print(extract(p), end="")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))

