from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


def _add_code_paragraph(doc: Document, s: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(s)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(10)


def convert(md_path: Path, out_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    in_code = False

    for raw in lines:
        line = raw.rstrip("\n")

        # Fenced code blocks (``` ... ```)
        if line.strip().startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            _add_code_paragraph(doc, line)
            continue

        if line.strip() == "---":
            doc.add_paragraph("")
            continue
        if not line.strip():
            doc.add_paragraph("")
            continue

        # Headings heuristics for this project style:
        # "1. ..." -> Heading 1
        # "1.1. ..." -> Heading 2
        if re.match(r"^\d+\.\d+\.", line.strip()):
            doc.add_paragraph(line.strip(), style="Heading 2")
            continue
        if re.match(r"^\d+\.\s", line.strip()):
            doc.add_paragraph(line.strip(), style="Heading 1")
            continue

        # Bullet list
        if line.lstrip().startswith("- "):
            doc.add_paragraph(line.lstrip()[2:], style="List Bullet")
            continue

        # Plain paragraph
        doc.add_paragraph(line)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main(argv: list[str]) -> int:
    sys.stdout.reconfigure(encoding="utf-8")
    if len(argv) != 3:
        print("Usage: python tools/md_to_docx.py <input.md> <output.docx>", file=sys.stderr)
        return 2

    md = Path(argv[1])
    out = Path(argv[2])
    if not md.exists():
        print(f"Input not found: {md}", file=sys.stderr)
        return 1

    convert(md, out)
    print(f"saved {out} size {out.stat().st_size}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))

